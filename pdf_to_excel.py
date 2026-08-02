import io
import zipfile
import pdfplumber
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
import streamlit as st


def show():
    # CSS - Custom styling with Yellow Download Buttons
    st.markdown(
        """
        <style>
        header[data-testid="stHeader"] {
            visibility: hidden !important;
            display: none !important;
        }
        .block-container {
            padding: 0.8rem 1rem !important;
            max-width: 100% !important;
        }
        div.row-widget.stVerticalBlock {
            gap: 0.05rem !important;
        }
        h3, h4 {
            margin: 0px !important;
            font-size: 1.15rem !important;
        }
        .stTextInput label, .stSelectbox label, .stFileUploader label {
            font-size: 1.05rem !important;
            font-weight: 700 !important;
            color: #0f172a !important;
        }
        /* Styling download buttons to Yellow */
        div.stDownloadButton > button {
            background-color: #FFC107 !important;
            color: #000000 !important;
            font-size: 1rem !important;
            font-weight: bold !important;
            padding: 6px 12px !important;
            min-height: 40px !important;
            border: 1px solid #E0A800 !important;
        }
        div.stDownloadButton > button:hover {
            background-color: #E0A800 !important;
            color: #000000 !important;
        }
        </style>
    """,
        unsafe_allow_html=True,
    )

    st.markdown("### 📄 Batch PDF Part-wise Annexure Generator (Up to 20 PDFs)")

    col_in1, col_in2 = st.columns(2)
    with col_in1:
        ac_name_input = st.text_input("Enter AC No. & Name", value="20, CHANDNI CHOWK")
    with col_in2:
        st.info("💡 Tip: You can select up to 20 PDF files at once.")

    uploaded_files = st.file_uploader(
        "Upload Voter PDF Documents (Select up to 20 files)", type=["pdf"], accept_multiple_files=True
    )

    if uploaded_files:
        if len(uploaded_files) > 20:
            st.warning("⚠️ Aapne 20 se zyada PDFs select ki hain. Kripya ek baar mein maximum 20 PDFs hi upload karein taaki process fast rahe.")
        else:
            st.success(f"📁 Total {len(uploaded_files)} PDF(s) Uploaded Successfully!")

            if st.button("📦 Generate & Download Batch ZIP (Excel & Word)", type="primary", use_container_width=True):
                with st.spinner("Processing your PDFs and generating part-wise Annexures..."):
                    
                    zip_buffer = io.BytesIO()
                    
                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                        
                        for uploaded_file in uploaded_files:
                            file_name_clean = uploaded_file.name.rsplit('.', 1)[0]
                            part_no = "".join([c for c in file_name_clean if c.isdigit()])
                            if not part_no:
                                part_no = file_name_clean

                            all_rows = []
                            with pdfplumber.open(uploaded_file) as pdf:
                                for page in pdf.pages:
                                    tables = page.extract_tables()
                                    for table in tables:
                                        for row in table:
                                            if row and any(row):
                                                cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                                row_str = "".join(cleaned_row).lower()
                                                
                                                if "epic" in row_str and "serial" in row_str:
                                                    continue
                                                if "elector name" in row_str:
                                                    continue
                                                all_rows.append(cleaned_row)

                            if all_rows:
                                clean_rows = []
                                for row in all_rows:
                                    row_str = "".join(row).lower()
                                    if not ("epic" in row_str and "serial" in row_str):
                                        clean_rows.append(row)

                                if clean_rows:
                                    df = pd.DataFrame(clean_rows)
                                    df = df.fillna("")

                                    if len(df.columns) >= 7:
                                        df = df.iloc[:, :7]
                                        df.columns = [
                                            "Original S.No", "Serial No", "EPIC Number", 
                                            "Elector Name", "Relative Details", "DOB/Age", "Uncollectable Reason"
                                        ]
                                    elif len(df.columns) == 6:
                                        df.columns = [
                                            "Serial No", "EPIC Number", "Elector Name", 
                                            "Relative Details", "DOB/Age", "Uncollectable Reason"
                                        ]
                                    else:
                                        while len(df.columns) < 7:
                                            df[f"Col_{len(df.columns)}"] = ""
                                        df.columns = [
                                            "Original S.No", "Serial No", "EPIC Number", 
                                            "Elector Name", "Relative Details", "DOB/Age", "Uncollectable Reason"
                                        ]

                                    df = df[
                                        (df["EPIC Number"].astype(str).str.strip() != "") | 
                                        (df["Elector Name"].astype(str).str.strip() != "")
                                    ]

                                    reason_col = "Uncollectable Reason"

                                    def get_annexure_df(category_key):
                                        filtered_df = df.copy()
                                        if category_key == "I":
                                            filtered_df = filtered_df[filtered_df[reason_col].astype(str).str.contains("absent|untraceable", case=False, na=False)]
                                        elif category_key == "II":
                                            filtered_df = filtered_df[filtered_df[reason_col].astype(str).str.contains("shifted|moved", case=False, na=False)]
                                        elif category_key == "III":
                                            filtered_df = filtered_df[filtered_df[reason_col].astype(str).str.contains("dead|expired|death", case=False, na=False)]
                                        elif category_key == "IV":
                                            filtered_df = filtered_df[filtered_df[reason_col].astype(str).str.contains("duplicate|already enrolled|repeat", case=False, na=False)]
                                        elif category_key == "V":
                                            filtered_df = filtered_df[filtered_df[reason_col].astype(str).str.contains("refused|ef|sign|other", case=False, na=False)]
                                        
                                        exp_df = pd.DataFrame()
                                        if not filtered_df.empty:
                                            exp_df["S. No."] = range(1, len(filtered_df) + 1)
                                            exp_df["EPIC No."] = filtered_df["EPIC Number"].values
                                            exp_df["SL. No. in the Part"] = filtered_df["Serial No"].values
                                            exp_df["Name of the Elector"] = filtered_df["Elector Name"].values
                                            exp_df["Name of Relative"] = filtered_df["Relative Details"].values
                                            exp_df["Category / Remarks"] = filtered_df[reason_col].values
                                        else:
                                            exp_df = pd.DataFrame(columns=["S. No.", "EPIC No.", "SL. No. in the Part", "Name of the Elector", "Name of Relative", "Category / Remarks"])
                                        return exp_df

                                    annexures_dict = {
                                        "I": ("List of ASDD Electors: Category- Absent", get_annexure_df("I")),
                                        "II": ("List of ASDD Electors: Category- Permanently Shifted", get_annexure_df("II")),
                                        "III": ("List of ASDD Electors: Category- Dead", get_annexure_df("III")),
                                        "IV": ("List of ASDD Electors: Category- Duplicate", get_annexure_df("IV")),
                                        "V": ("List of ASDD Electors: Category- Others", get_annexure_df("V"))
                                    }

                                    # 1. Generate Excel for this Part
                                    excel_output = io.BytesIO()
                                    with pd.ExcelWriter(excel_output, engine='openpyxl') as writer:
                                        for key, (title_desc, ann_df) in annexures_dict.items():
                                            if ann_df.empty:
                                                ann_df = pd.DataFrame(columns=["S. No.", "EPIC No.", "SL. No. in the Part", "Name of the Elector", "Name of Relative", "Category / Remarks"])
                                            ann_df.to_excel(writer, index=False, sheet_name=f"Annexure-{key}", startrow=3)

                                    excel_output.seek(0)
                                    wb = openpyxl.load_workbook(excel_output)
                                    
                                    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
                                    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
                                    data_font = Font(name="Calibri", size=10)
                                    thin_border = Border(
                                        left=Side(style='thin', color='D9D9D9'), right=Side(style='thin', color='D9D9D9'),
                                        top=Side(style='thin', color='D9D9D9'), bottom=Side(style='thin', color='D9D9D9')
                                    )

                                    for key, (title_desc, ann_df) in annexures_dict.items():
                                        ws = wb[f"Annexure-{key}"]
                                        ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
                                        ws.page_setup.paperSize = ws.PAPERSIZE_A4
                                        ws.sheet_properties.pageSetUpPr.fitToPage = True
                                        ws.page_setup.fitToWidth = 1
                                        ws.page_setup.fitToHeight = 0

                                        ws['A1'] = f"AC No. & Name: {ac_name_input}"
                                        cell_ann = ws['F1']
                                        cell_ann.value = f"Annexure- {key}"
                                        cell_ann.font = Font(name="Calibri", size=11, bold=True)
                                        cell_ann.alignment = Alignment(horizontal="right", vertical="center")

                                        ws['A2'] = f"Part No. {part_no}"
                                        ws.merge_cells(start_row=3, start_column=1, end_row=3, end_column=6)
                                        cell_title = ws['A3']
                                        cell_title.value = title_desc
                                        cell_title.font = Font(name="Calibri", size=12, bold=True)
                                        cell_title.alignment = Alignment(horizontal="center", vertical="center")

                                        table_header_row = 4
                                        for col_idx in range(1, ws.max_column + 1):
                                            cell = ws.cell(row=table_header_row, column=col_idx)
                                            cell.fill = header_fill
                                            cell.font = header_font
                                            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                                            cell.border = thin_border

                                        for row in range(table_header_row + 1, ws.max_row + 1):
                                            for col in range(1, ws.max_column + 1):
                                                cell = ws.cell(row=row, column=col)
                                                cell.font = data_font
                                                cell.border = thin_border
                                                if col in [1, 2, 3]:
                                                    cell.alignment = Alignment(horizontal="center", vertical="center")
                                                else:
                                                    cell.alignment = Alignment(horizontal="left", vertical="center")

                                        ws.column_dimensions['A'].width = 8
                                        ws.column_dimensions['B'].width = 22
                                        ws.column_dimensions['C'].width = 20
                                        ws.column_dimensions['D'].width = 32
                                        ws.column_dimensions['E'].width = 35
                                        ws.column_dimensions['F'].width = 25

                                    part_excel_io = io.BytesIO()
                                    wb.save(part_excel_io)
                                    zip_file.writestr(f"Excel_Files/Part_No_{part_no}_All_Annexures.xlsx", part_excel_io.getvalue())

                                    # 2. Generate Word for this Part
                                    doc = docx.Document()
                                    for section in doc.sections:
                                        section.top_margin = Inches(0.8)
                                        section.bottom_margin = Inches(0.8)
                                        section.left_margin = Inches(0.8)
                                        section.right_margin = Inches(0.8)

                                    for key, (title_desc, ann_df) in annexures_dict.items():
                                        p_top = doc.add_paragraph()
                                        r_top1 = p_top.add_run(f"AC No. & Name: {ac_name_input}")
                                        r_top1.font.name = 'Calibri'
                                        r_top1.font.size = Pt(10.5)
                                        r_top1.font.bold = True

                                        p_top.add_run("\t\t\t\t\t\t")
                                        r_top_ann = p_top.add_run(f"Annexure- {key}")
                                        r_top_ann.font.name = 'Calibri'
                                        r_top_ann.font.size = Pt(11)
                                        r_top_ann.font.bold = True

                                        p_part = doc.add_paragraph()
                                        r_part = p_part.add_run(f"Part No. {part_no}\n")
                                        r_part.font.name = 'Calibri'
                                        r_part.font.size = Pt(10.5)
                                        r_part.font.bold = True

                                        p_title = doc.add_paragraph()
                                        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        r_title = p_title.add_run(f"{title_desc}")
                                        r_title.font.name = 'Calibri'
                                        r_title.font.size = Pt(11.5)
                                        r_title.font.bold = True

                                        doc.add_paragraph()

                                        if ann_df.empty:
                                            p_empty = doc.add_paragraph(f"No records found for {title_desc}")
                                            p_empty.runs[0].font.italic = True
                                        else:
                                            table = doc.add_table(rows=len(ann_df) + 1, cols=len(ann_df.columns))
                                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                            table.style = 'Table Grid'

                                            col_widths = [Inches(0.6), Inches(1.4), Inches(1.2), Inches(2.2), Inches(2.2), Inches(1.6)]
                                            for row in table.rows:
                                                for idx, width in enumerate(col_widths):
                                                    row.cells[idx].width = width

                                            hdr_cells = table.rows[0].cells
                                            for i, col_name in enumerate(ann_df.columns):
                                                hdr_cells[i].text = str(col_name)
                                                for paragraph in hdr_cells[i].paragraphs:
                                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                    for run in paragraph.runs:
                                                        run.font.bold = True
                                                        run.font.name = 'Calibri'
                                                        run.font.size = Pt(10)
                                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                                shading_elm = parse_xml(r'<w:shd {} w:fill="1F4E78"/>'.format(docx.oxml.ns.nsdecls('w')))
                                                hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

                                            for row_idx, row in ann_df.iterrows():
                                                row_cells = table.rows[row_idx + 1].cells
                                                for col_idx, val in enumerate(row):
                                                    row_cells[col_idx].text = str(val if val is not None else "")
                                                    for paragraph in row_cells[col_idx].paragraphs:
                                                        for run in paragraph.runs:
                                                            run.font.name = 'Calibri'
                                                            run.font.size = Pt(10)

                                        p_sign = doc.add_paragraph("\n")
                                        p_sign.add_run("Sign of BLO _________________________                                           Sign of BLA-2 (i) ____________________\n")
                                        p_sign.add_run("                                                                                                                            (ii) ____________________\n")
                                        p_sign.add_run("                                                                                                                           (iii) ____________________\n")
                                        p_sign.add_run("                                                                                                                           (iv) ____________________\n\n")
                                        for run in p_sign.runs:
                                            run.font.name = 'Calibri'
                                            run.font.size = Pt(9.5)
                                        
                                        doc.add_page_break()

                                    part_word_io = io.BytesIO()
                                    doc.save(part_word_io)
                                    zip_file.writestr(f"Word_Files/Part_No_{part_no}_All_Annexures.docx", part_word_io.getvalue())

                    zip_buffer.seek(0)
                    st.success("✅ Batch Process Complete! All files are packed in ZIP.")
                    
                    st.download_button(
                        label="📥 Download Batch Annexures (ZIP)",
                        data=zip_buffer.getvalue(),
                        file_name="Batch_Parts_Annexures_I_to_V.zip",
                        mime="application/zip",
                        use_container_width=True,
                    )
    else:
        st.info("👉 Please upload up to 20 Voter PDF files to start conversion.")