from datetime import datetime
import io
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pypdf
import re
import streamlit as st


def process_voter_pdf(pdf_file):
  """पीडीएफ से डेटा छानकर 8वें नंबर के रीज़न के आधार पर सही एनेवेक्सचर में बांटना"""
  reader = pypdf.PdfReader(pdf_file)

  ac_info = "20 - CHANDNI CHOWK"
  part_info = "1-MATIA MAHAL"
  part_filename_str = "Part_1"

  if len(reader.pages) > 0:
    first_text = reader.pages[0].extract_text()
    for line in first_text.split("\n"):
      if "part" in line.lower():
        part_info = line.strip()
        nums = re.findall(r"\d+", line)
        if len(nums) >= 2:
          part_filename_str = f"Part_{nums[1]}"
        elif len(nums) == 1:
          part_filename_str = f"Part_{nums[0]}"
        break

  annexure_1 = []  # Absent / Untraceable
  annexure_2 = []  # Permanently Shifted
  annexure_3 = []  # Dead
  annexure_4 = []  # Duplicate -> Already Enrolled
  annexure_5 = []  # Others -> EF Refused

  for page in reader.pages:
    text = page.extract_text()
    lines = [l.strip() for l in text.split("\n") if l.strip()]

    i = 0
    while i < len(lines):
      line = lines[i]
      line_lower = line.lower()

      if (
          not line
          or len(line) < 3
          or "page" in line_lower
          or "generated" in line_lower
          or "note" in line_lower
          or "list" in line_lower
          or "electoral" in line_lower
          or "assembly" in line_lower
          or "ilovepdf" in line_lower
      ):
        i += 1
        continue

      if "|" in line or (line[0].isdigit() and not "ac no" in line_lower):
        block_lines = [line]
        j = i + 1
        while j < len(lines) and j < i + 4:
          nxt = lines[j]
          if "|" in nxt or (nxt and nxt[0].isdigit() and len(nxt) < 4):
            break
          if nxt:
            block_lines.append(nxt)
          j += 1

        full_block = " ".join(block_lines)
        full_block_lower = full_block.lower()
        tokens = full_block.split()

        epic_match = re.search(r"\b[A-Z]{3}\d{7}\b", full_block)
        if not epic_match:
          i += 1
          continue
        epic_val = epic_match.group(0)

        sl_no_val = tokens[1] if len(tokens) > 1 and tokens[1].isdigit() else "1"
        elector_name = tokens[3].upper() if len(tokens) > 3 else "N/A"
        relative_name = tokens[4].upper() if len(tokens) > 4 else "N/A"

        reason = ""
        target_annexure = 5

        token_8 = tokens[7].lower() if len(tokens) > 7 else ""

        if (
            "death" in token_8
            or "dead" in token_8
            or "death" in full_block_lower
            or "dead" in full_block_lower
        ):
          reason = "Dead"
          target_annexure = 3
        elif (
            "shift" in token_8
            or "permanently" in token_8
            or "shift" in full_block_lower
            or "permanently" in full_block_lower
        ):
          reason = "Permanently Shifted"
          target_annexure = 2
        elif (
            "already" in token_8
            or "enrolled" in token_8
            or "duplicate" in token_8
            or "enroll" in token_8
            or "already enrolled" in full_block_lower
            or "duplicate" in full_block_lower
        ):
          reason = "Already Enrolled"
          target_annexure = 4
        elif (
            "untraceable" in token_8
            or "absent" in token_8
            or "untraceable" in full_block_lower
            or "absent" in full_block_lower
        ):
          reason = "Untraceable/Absent"
          target_annexure = 1
        elif (
            "ef" in token_8
            or "refused" in token_8
            or "ef refused" in full_block_lower
        ):
          reason = "EF Refused"
          target_annexure = 5
        else:
          i += 1
          continue

        row_data = {
            "EPIC No.": epic_val,
            "SL. No. in the Part": sl_no_val,
            "Name of the Elector": elector_name,
            "Name of Relative": relative_name,
            "Remarks": reason,
        }

        if target_annexure == 3:
          annexure_3.append(row_data)
        elif target_annexure == 2:
          annexure_2.append(row_data)
        elif target_annexure == 4:
          annexure_4.append(row_data)
        elif target_annexure == 1:
          annexure_1.append(row_data)
        else:
          annexure_5.append(row_data)

      i += 1

  return (
      ac_info,
      part_info,
      part_filename_str,
      annexure_1,
      annexure_2,
      annexure_3,
      annexure_4,
      annexure_5,
  )


def add_minutes_of_meeting_document(
    doc, part_info, count_a1, count_a2, count_a3, count_a4, count_a5
):
  """डॉक्यूमेंट के अंत में मिनट्स ऑफ़ मीटिंग जोड़ने का फंक्शन"""
  doc.add_page_break()

  p_head = doc.add_paragraph()
  p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run_h1 = p_head.add_run(
      "OFFICE OF THE ASSISTANT ELECTORAL REGISTRATION OFFICER\n"
  )
  run_h1.bold = True
  run_h1.font.size = Pt(11)
  run_h2 = p_head.add_run(
      "ASSEMBLY CONSTITUENCY NUMBER 20, & NAME: Chandni Chowk\n\n"
  )
  run_h2.bold = True
  run_h2.font.size = Pt(10)

  p_title = doc.add_paragraph()
  p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
  run_title = p_title.add_run("MINUTES OF MEETING\n")
  run_title.bold = True
  run_title.font.size = Pt(12)
  p_title.paragraph_format.space_after = Pt(10)

  p1 = doc.add_paragraph()
  p1.add_run(
      "In pursuance of the directions of the Election Commission of India, with"
      " regard to sharing of the information in respect of Uncollectable"
      " Enumeration Forms on account of different reasons, a meeting with the"
      " Booth Level Agents-2 (BLAs-2) of recognized political parties was held"
      " today on (Date) ____________ at ____________ (Time) at the (Venue)"
      " ________________________.\n"
  )
  p1.runs[0].font.size = Pt(10)

  p2 = doc.add_paragraph()
  p2.add_run("2. Following attended the meeting:-\n")
  p2.runs[0].bold = True
  p2.runs[0].font.size = Pt(10)

  table_att = doc.add_table(rows=6, cols=3)
  table_att.alignment = WD_TABLE_ALIGNMENT.CENTER
  table_att.style = "Table Grid"

  headers_att = ["S. No.", "Name", "Name of the Political party"]
  for i, h in enumerate(headers_att):
    table_att.rows[0].cells[i].text = h
    table_att.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    table_att.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

  for row_idx in range(1, 6):
    row_cells = table_att.rows[row_idx].cells
    row_cells[0].text = str(row_idx)
    row_cells[1].text = ""
    row_cells[2].text = ""
    for cell in row_cells:
      for p in cell.paragraphs:
        for r in p.runs:
          r.font.size = Pt(10)

  p_space = doc.add_paragraph()
  p_space.paragraph_format.space_after = Pt(6)

  p3 = doc.add_paragraph()
  p3.add_run(
      "3. At the outset, all the participants were welcomed by the"
      " undersigned.\n"
  )
  p3.runs[0].font.size = Pt(10)

  p4 = doc.add_paragraph()
  p4.add_run(
      f"4. It was informed that a total of ________ electors existed prior to"
      f" start of SIR in Part No. {part_info}.\n"
  )
  p4.runs[0].font.size = Pt(10)

  total_electors = count_a1 + count_a2 + count_a3 + count_a4 + count_a5
  p5 = doc.add_paragraph()
  p5.add_run(
      f"5. Out of total electors, {total_electors} no. of Enumeration Forms"
      " have been collected during House to House visits.\n"
  )
  p5.runs[0].font.size = Pt(10)

  p6 = doc.add_paragraph()
  p6.add_run(
      "6. Based on the House to House visits, a list of uncollectable"
      " Enumeration Forms was prepared as per instructions of the ECI. The same"
      " was shared with the BLAs as per details given below:-\n"
  )
  p6.runs[0].font.size = Pt(10)

  table_sum = doc.add_table(rows=6, cols=3)
  table_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
  table_sum.style = "Table Grid"

  headers_sum = ["S. No.", "Types of Electors", "Number of Such Electors"]
  for i, h in enumerate(headers_sum):
    table_sum.rows[0].cells[i].text = h
    table_sum.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    table_sum.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)

  sum_data = [
      ("1", "Absent", str(count_a1)),
      ("2", "Shifted", str(count_a2)),
      ("3", "Dead", str(count_a3)),
      ("4", "Duplicate", str(count_a4)),
      ("5", "Others (refused to sign EF, Demolished etc )", str(count_a5)),
  ]

  for r_idx, (sno, etype, evalue) in enumerate(sum_data, start=1):
    row_cells = table_sum.rows[r_idx].cells
    row_cells[0].text = sno
    row_cells[1].text = etype
    row_cells[2].text = evalue
    for cell in row_cells:
      for p in cell.paragraphs:
        for r in p.runs:
          r.font.size = Pt(10)

  doc.add_paragraph("\n")

  p7 = doc.add_paragraph()
  p7.add_run(
      "7. All BLAs were satisfied with the list of uncollectable Enumeration"
      " Forms provided to them.\n"
  )
  p7.runs[0].font.size = Pt(10)

  p8 = doc.add_paragraph()
  p8.add_run(
      "8. It was also informed to the BLAs present in the meeting that"
      " objections if any, may be filed during the period of claims & objection"
      " period i.e. 05.08.2026 to 04.09.2026.\n"
  )
  p8.runs[0].font.size = Pt(10)

  p9 = doc.add_paragraph()
  p9.add_run("9. Meeting ended with vote of thanks.\n\n")
  p9.runs[0].font.size = Pt(10)

  p_sign = doc.add_paragraph()
  p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
  run_s = p_sign.add_run(
      "Signatures of the BLO\n\n(Name of BLO) ______________\nPS Number"
      " ______________"
  )
  run_s.font.size = Pt(10)

  p_bla = doc.add_paragraph()
  p_bla.add_run("Signature of BLAs (i) ____________________\n")
  p_bla.add_run("                   (ii) ____________________\n")
  p_bla.add_run("                  (iii) ____________________\n")
  p_bla.add_run("                  (iv) ____________________\n\n")
  for r in p_bla.runs:
    r.font.size = Pt(10)

  p_end = doc.add_paragraph()
  p_end.add_run(
      "Encls:\n1. Lists of Uncollectable Enumeration Forms/ASDD List.\n2."
      " Photographs of the meeting.\n\nCopy to:-\n1. AERO AC Number & Name: 20,"
      " Chandni Chowk\n2. The BLA ________________________"
  )
  p_end.runs[0].font.size = Pt(9)


def generate_full_report_document(ac_info, part_info, a1, a2, a3, a4, a5):
  """एनेवेक्सचर रिपोर्ट और अंत में मिनट्स ऑफ़ मीटिंग जोड़ने वाला मुख्य फंक्शन"""
  doc = docx.Document()

  annexure_configs = [
      ("Annexure- I", "List of ASDD Electors: Category- Absent", a1),
      (
          "Annexure- II",
          "List of ASDD Electors: Category- Permanently Shifted",
          a2,
      ),
      ("Annexure- III", "List of ASDD Electors: Category- Dead", a3),
      ("Annexure- IV", "List of ASDD Electors: Category- Duplicate", a4),
      ("Annexure- V", "List of ASDD Electors: Category- Others", a5),
  ]

  for idx, (ann_no, ann_title, data) in enumerate(annexure_configs):
    p_ann = doc.add_paragraph()
    p_ann.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_ann = p_ann.add_run(ann_no)
    run_ann.bold = True
    run_ann.font.size = Pt(10)

    p_ac = doc.add_paragraph(
        "Assembly Constituency Number 20, & Name: Chandni Chowk"
    )
    p_ac.runs[0].font.size = Pt(10)
    p_part = doc.add_paragraph(f"Part No.: {part_info}")
    p_part.runs[0].font.size = Pt(10)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(ann_title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    p_title.paragraph_format.space_after = Pt(10)

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    headers = [
        "S. No.",
        "EPIC No.",
        "SL. No. in the Part",
        "Name of the Elector",
        "Name of Relative",
        "Remarks",
    ]

    for i, h_text in enumerate(headers):
      hdr_cells[i].text = h_text
      p = hdr_cells[i].paragraphs[0]
      p.alignment = WD_ALIGN_PARAGRAPH.CENTER
      if len(p.runs) > 0:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10)

    rows_to_add = data[:200] if data else []

    if not rows_to_add:
      rows_to_add = [{
          "EPIC No.": "N/A",
          "SL. No. in the Part": "",
          "Name of the Elector": "No records found",
          "Name of Relative": "",
          "Remarks": "",
      }]

    for row_idx, row in enumerate(rows_to_add):
      row_cells = table.add_row().cells
      row_cells[0].text = str(row_idx + 1)
      row_cells[1].text = str(row.get("EPIC No.", ""))
      row_cells[2].text = str(row.get("SL. No. in the Part", ""))
      row_cells[3].text = str(row.get("Name of the Elector", ""))
      row_cells[4].text = str(row.get("Name of Relative", ""))
      row_cells[5].text = str(row.get("Remarks", ""))

      for cell in row_cells:
        for paragraph in cell.paragraphs:
          paragraph.paragraph_format.space_after = Pt(2)
          paragraph.paragraph_format.space_before = Pt(2)
          for run in paragraph.runs:
            run.font.size = Pt(10)

    doc.add_paragraph("\n")
    p_sign = doc.add_paragraph()
    p_sign.add_run("Sign of BLO ____________________").font.size = Pt(10)

    p_bla = doc.add_paragraph()
    p_bla.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_bla.add_run("Sign of BLA-2 (i) ____________________\n").font.size = Pt(10)
    p_bla.add_run("                 (ii) ____________________\n").font.size = (
        Pt(10)
    )
    p_bla.add_run("                (iii) ____________________\n").font.size = (
        Pt(10)
    )
    p_bla.add_run("                (iv) ____________________").font.size = Pt(
        10
    )

    doc.add_page_break()

  add_minutes_of_meeting_document(
      doc,
      part_info,
      len(a1),
      len(a2),
      len(a3),
      len(a4),
      len(a5),
  )

  doc_io = io.BytesIO()
  doc.save(doc_io)
  doc_io.seek(0)
  return doc_io


def show():
  st.markdown(
      """
        <style>
        header {visibility: hidden;}
        .block-container {
            padding-top: 0rem !important;
            padding-bottom: 2rem !important;
            margin-top: -20px !important;
        }
        h1 {
            color: #ff4b4b;
            font-weight: 800;
        }
        h3 {
            color: #4e73df;
        }
        .stDownloadButton > button {
            background-color: #28a745 !important;
            color: white !important;
            border-radius: 10px;
            font-weight: bold;
            border: none;
            box-shadow: 0 4px 10px rgba(40, 167, 69, 0.3);
            width: 100%;
        }
        .stDownloadButton > button:hover {
            background-color: #218838 !important;
        }
        </style>
    """,
      unsafe_allow_html=True,
  )

  current_time = datetime.now().strftime("%d %B %Y, %I:%M:%S %p")
  st.title("📄 Complete Voter Report Generator")
  st.caption(
      "Annexures + Minutes of Meeting on Last Page | 🕒 "
      f"{current_time}"
  )

  st.divider()

  uploaded_pdf = st.file_uploader(
      "Upload Voter PDF Report", type=["pdf"], key="exact_word_up_full_report_v7"
  )

  st.divider()

  if st.button(
      "🚀 Generate Complete Word Report", use_container_width=True, type="primary"
  ):
    if uploaded_pdf is None:
      st.warning("⚠️ कृपया पहले अपनी वोटर PDF फाइल अपलोड करें।")
    else:
      with st.spinner("⏳ एनेवेक्सचर और मिनट्स ऑफ़ मीटिंग के साथ रिपोर्ट बन रही है..."):
        try:
          (
              ac_info,
              part_info,
              part_filename_str,
              a1,
              a2,
              a3,
              a4,
              a5,
          ) = process_voter_pdf(uploaded_pdf)
          word_io = generate_full_report_document(
              ac_info, part_info, a1, a2, a3, a4, a5
          )

          st.success("🎉 आपकी पूरी वर्ड रिपोर्ट (एनेवेक्सचर + मिनट्स) तैयार है!")

          st.download_button(
              label=f"📥 Download {part_filename_str}_Complete_Report.docx",
              data=word_io,
              file_name=f"{part_filename_str}_Complete_Report.docx",
              mime=(
                  "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
              ),
          )
        except Exception as e:
          st.error(f"❌ त्रुटि: {e}")


if __name__ == "__main__":
  show()